import docx

def read_docx(filepath):
    doc = docx.Document(filepath)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    # Check tables too, just in case they have content
    tableText = []
    for table in doc.tables:
        for row in table.rows:
            rowText = [cell.text.strip() for cell in row.cells]
            tableText.append(" | ".join(rowText))
    return "\n".join(fullText), "\n".join(tableText)

p_text, t_text = read_docx(r'_posts\데이터베이스 기초 정리.docx')
print("--- 데이터베이스 기초 정리.docx Paragraphs ---")
print(p_text[:3000])
print("\n--- 데이터베이스 기초 정리.docx Tables ---")
print(t_text[:2000])

p_text2, t_text2 = read_docx(r'_posts\반도체장비데이터관리 정리.docx')
print("\n=== 반도체장비데이터관리 정리.docx Paragraphs ===")
print(p_text2[:3000])
print("\n=== 반도체장비데이터관리 정리.docx Tables ===")
print(t_text2[:2000])
